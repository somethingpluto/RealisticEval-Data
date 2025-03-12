f
r
o
m
 
d
o
c
x
 
i
m
p
o
r
t
 
D
o
c
u
m
e
n
t




d
e
f
 
e
x
t
r
a
c
t
_
t
e
x
t
_
f
r
o
m
_
w
o
r
d
(
d
o
c
x
_
f
i
l
e
_
p
a
t
h
)
:


 
 
 
 
"
"
"


 
 
 
 
E
x
t
r
a
c
t
s
 
t
e
x
t
 
c
o
n
t
e
n
t
 
f
r
o
m
 
a
 
g
i
v
e
n
 
W
o
r
d
 
f
i
l
e
 
(
.
d
o
c
x
)
.




 
 
 
 
A
r
g
s
:


 
 
 
 
 
 
 
 
d
o
c
x
_
f
i
l
e
_
p
a
t
h
 
(
s
t
r
)
:
 
T
h
e
 
p
a
t
h
 
t
o
 
t
h
e
 
W
o
r
d
 
f
i
l
e
.




 
 
 
 
R
e
t
u
r
n
s
:


 
 
 
 
 
 
 
 
s
t
r
:
 
T
h
e
 
e
x
t
r
a
c
t
e
d
 
t
e
x
t
 
c
o
n
t
e
n
t
.


 
 
 
 
"
"
"


 
 
 
 
d
o
c
u
m
e
n
t
 
=
 
D
o
c
u
m
e
n
t
(
d
o
c
x
_
f
i
l
e
_
p
a
t
h
)


 
 
 
 
t
e
x
t
 
=
 
'
'


 
 
 
 
f
o
r
 
p
a
r
a
g
r
a
p
h
 
i
n
 
d
o
c
u
m
e
n
t
.
p
a
r
a
g
r
a
p
h
s
:


 
 
 
 
 
 
 
 
t
e
x
t
 
+
=
 
p
a
r
a
g
r
a
p
h
.
t
e
x
t
 
+
 
'
\
n
'


 
 
 
 
r
e
t
u
r
n
 
t
e
x
t




d
e
f
 
e
x
t
r
a
c
t
_
t
e
x
t
_
f
r
o
m
_
p
d
f
(
p
d
f
_
f
i
l
e
_
p
a
t
h
)
:


 
 
 
 
"
"
"


 
 
 
 
E
x
t
r
a
c
t
s
 
t
e
x
t
 
c
o
n
t
e
n
t
 
f
r
o
m
 
a
 
g
i
v
e
n
 
P
D
F
 
f
i
l
e
 
(
.
p
d
f
)
.




 
 
 
 
A
r
g
s
:


 
 
 
 
 
 
 
 
p
d
f
_
f
i
l
e
_
p
a
t
h
 
(
s
t
r
)
:
 
T
h
e
 
p
a
t
h
 
t
o
 
t
h
e
 
P
D
F
 
f
i
l
e
.




 
 
 
 
R
e
t
u
r
n
s
:


 
 
 
 
 
 
 
 
s
t
r
:
 
T
h
e
 
e
x
t
r
a
c
t
e
d
 
t
e
x
t
 
c
o
n
t
e
n
t
.


 
 
 
 
"
"
"


 
 
 
 
i
m
p
o
r
t
 
o
s


 
 
 
 
i
m
p
o
r
t
 
s
y
s


 
 
 
 
i
m
p
o
r
t
 
s
u
b
p
r
o
c
e
s
s


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
p
d
f
i
n
t
e
r
p
 
i
m
p
o
r
t
 
P
D
F
R
e
s
o
u
r
c
e
M
a
n
a
g
e
r
,
 
P
D
F
P
a
g
e
I
n
t
e
r
p
r
e
t
e
r


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
c
o
n
v
e
r
t
e
r
 
i
m
p
o
r
t
 
T
e
x
t
C
o
n
v
e
r
t
e
r


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
l
a
y
o
u
t
 
i
m
p
o
r
t
 
L
A
P
a
r
a
m
s


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
p
d
f
p
a
g
e
 
i
m
p
o
r
t
 
P
D
F
P
a
g
e


 
 
 
 
f
r
o
m
 
i
o
 
i
m
p
o
r
t
 
S
t
r
i
n
g
I
O


 
 
 
 
f
r
o
m
 
i
o
 
i
m
p
o
r
t
 
o
p
e
n


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
p
d
f
p
a
r
s
e
r
 
i
m
p
o
r
t
 
P
D
F
P
a
r
s
e
r


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
p
d
f
d
o
c
u
m
e
n
t
 
i
m
p
o
r
t
 
P
D
F
D
o
c
u
m
e
n
t


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
p
d
f
p
a
g
e
 
i
m
p
o
r
t
 
P
D
F
P
a
g
e


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
p
d
f
p
a
g
e
 
i
m
p
o
r
t
 
P
D
F
T
e
x
t
E
x
t
r
a
c
t
i
o
n
N
o
t
A
l
l
o
w
e
d


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
p
d
f
i
n
t
e
r
p
 
i
m
p
o
r
t
 
P
D
F
R
e
s
o
u
r
c
e
M
a
n
a
g
e
r


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
p
d
f
i
n
t
e
r
p
 
i
m
p
o
r
t
 
P
D
F
P
a
g
e
I
n
t
e
r
p
r
e
t
e
r


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
p
d
f
d
e
v
i
c
e
 
i
m
p
o
r
t
 
P
D
F
D
e
v
i
c
e


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
l
a
y
o
u
t
 
i
m
p
o
r
t
 
L
A
P
a
r
a
m
s


 
 
 
 
f
r
o
m
 
p
d
f
m
i
n
e
r
.
c
o
n
v
e
r
t
e
r
 
i
m
p
o
r
t
 
P
D
F
P
a
g
e
A
g
g
r
e
g
a
t
o
r




 
 
 
 
#
 
O
p
e
n
 
a
 
P
D
F
 
f
i
l
e
.


 
 
 
 
f
p
 
=
 
o
p
e
n
(
p
d
f
_
f
i
l
e
_
p
a
t
h
,
 
'
r
b
'
)




 
 
 
 
#
 
C
r
e
a
t
e
 
a
 
P
D
F
 
p
a
r
s
e
r
 
o
b
j
e
c
t
 
a
s
s
o
c
i
a
t
e
d
 
w
i
t
h
 
t
h
e
 
f
i
l
e
 
o
b
j
e
c
t
.


 
 
 
 
p
a
r
s
e
r
 
=
 
P
D
F
P
a
r
s
e
r
(
f
p
)




 
 
 
 
#
 
C
r
e
a
t
e
 
a
 
P
D
F
 
d
o
c
u
m
e
n
t
 
o
b
j
e
c
t
 
t
h
a
t
 
s
t
o
r
e
s
 
t
h
e
 
d
o
c
u
m
e
n
t
 
s
t
r
u
c
t
u
r
e
.


 
 
 
 
#
 
P
a
s
s
w
o
r
d
 
f
o
r
 
i
n
i
t
i
a
l
i
z
a
t
i
o
n
 
a
s
 
2
n
d
 
p
a
r
a
m
e
t
e
r


 
 
 
 
d
o
c
u
m
e
n
t
 
=
 
P
D
F
D
o
c
u
m
e
n
t
(
p
a
r
s
e
r
)




 
 
 
 
#
 
C
h
e
c
k
 
i
f
 
t
h
e
 
d
o
c
u
m
e
n
t
 
a
l
l
o
w
s
 
t
e
x
t
 
e
x
t
r
a
c
t
i
o
n
.
 
I
f
 
n
o
t
,
 
a
b
o
r
t
.


 
 
 
 
i
f
import unittest
import os
from docx import Document


class TestExtractTextFromWord(unittest.TestCase):

    def setUp(self):
        """Set up the testing environment."""
        # Create a temporary Word file for testing
        self.test_docx_path = "test_document.docx"
        self.create_sample_docx()

    def tearDown(self):
        """Clean up the test environment."""
        # Remove created files after tests
        if os.path.exists(self.test_docx_path):
            os.remove(self.test_docx_path)

    def create_sample_docx(self):
        """Helper method to create a sample Word document for testing."""
        doc = Document()
        doc.add_paragraph("Hello World!")
        doc.add_paragraph("This is a test document.")
        doc.save(self.test_docx_path)

    def test_extract_text_success(self):
        """Test extracting text from a normal Word document."""
        expected_text = "Hello World!\nThis is a test document."
        extracted_text = extract_text_from_word(self.test_docx_path)
        self.assertEqual(extracted_text.strip(), expected_text)

    def test_extract_empty_document(self):
        """Test extracting text from an empty Word document."""
        empty_docx_path = "empty_document.docx"
        Document().save(empty_docx_path)

        extracted_text = extract_text_from_word(empty_docx_path)
        self.assertEqual(extracted_text, "")  # Expecting an empty string

        os.remove(empty_docx_path)  # Clean up


    def test_extract_text_with_special_characters(self):
        """Test extracting text from a document containing special characters."""
        special_docx_path = "special_characters.docx"
        doc = Document()
        doc.add_paragraph("Hello, 世界! @#$%^&*()")
        doc.save(special_docx_path)

        extracted_text = extract_text_from_word(special_docx_path)
        expected_text = "Hello, 世界! @#$%^&*()"
        self.assertEqual(extracted_text.strip(), expected_text)

        os.remove(special_docx_path)  # Clean up

    def test_extract_text_with_multiple_paragraphs(self):
        """Test extracting text from a document with multiple paragraphs."""
        multi_para_docx_path = "multi_paragraphs.docx"
        doc = Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        doc.add_paragraph("Third paragraph.")
        doc.save(multi_para_docx_path)

        extracted_text = extract_text_from_word(multi_para_docx_path)
        expected_text = "First paragraph.\nSecond paragraph.\nThird paragraph."
        self.assertEqual(extracted_text.strip(), expected_text)

        os.remove(multi_para_docx_path)  # Clean up

if __name__ == '__main__':
    unittest.main()