import unittest
import os
from docx import Document


class TestExtractTextFromWord(unittest.TestCase):

    def setUp(self):
        """Set up the testing environment."""
        # Create a temporary Word file for testing
        self.test_docx_path = "test_document.docx"
        self.create_sample_docx()

    def tearDown(self):
        """Clean up the test environment."""
        # Remove created files after tests
        if os.path.exists(self.test_docx_path):
            os.remove(self.test_docx_path)

    def create_sample_docx(self):
        """Helper method to create a sample Word document for testing."""
        doc = Document()
        doc.add_paragraph("Hello World!")
        doc.add_paragraph("This is a test document.")
        doc.save(self.test_docx_path)

    def test_extract_text_success(self):
        """Test extracting text from a normal Word document."""
        expected_text = "Hello World!\nThis is a test document."
        extracted_text = extract_text_from_word(self.test_docx_path)
        self.assertEqual(extracted_text.strip(), expected_text)

    def test_extract_empty_document(self):
        """Test extracting text from an empty Word document."""
        empty_docx_path = "empty_document.docx"
        Document().save(empty_docx_path)

        extracted_text = extract_text_from_word(empty_docx_path)
        self.assertEqual(extracted_text, "")  # Expecting an empty string

        os.remove(empty_docx_path)  # Clean up


    def test_extract_text_with_special_characters(self):
        """Test extracting text from a document containing special characters."""
        special_docx_path = "special_characters.docx"
        doc = Document()
        doc.add_paragraph("Hello, 世界! @#$%^&*()")
        doc.save(special_docx_path)

        extracted_text = extract_text_from_word(special_docx_path)
        expected_text = "Hello, 世界! @#$%^&*()"
        self.assertEqual(extracted_text.strip(), expected_text)

        os.remove(special_docx_path)  # Clean up

    def test_extract_text_with_multiple_paragraphs(self):
        """Test extracting text from a document with multiple paragraphs."""
        multi_para_docx_path = "multi_paragraphs.docx"
        doc = Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        doc.add_paragraph("Third paragraph.")
        doc.save(multi_para_docx_path)

        extracted_text = extract_text_from_word(multi_para_docx_path)
        expected_text = "First paragraph.\nSecond paragraph.\nThird paragraph."
        self.assertEqual(extracted_text.strip(), expected_text)

        os.remove(multi_para_docx_path)  # Clean up
